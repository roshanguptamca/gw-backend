import io
import json
import logging
import re
from calendar import monthrange
from collections import Counter
from datetime import date, timedelta

from django.conf import settings
from django.core.serializers.json import DjangoJSONEncoder
from django.db import transaction
from django.forms.models import model_to_dict
from django.utils import timezone

import pdfplumber
from docx import Document as DocxDocument
from pypdf import PdfReader

from apps.resumes.models import (
    Award,
    Certification,
    Education,
    PersonalDetail,
    Project,
    Reference,
    ResumeLanguage,
    ResumeSummary,
    Skill,
    TemporaryResumeUpload,
    WorkExperience,
)

logger = logging.getLogger(__name__)
RESUME_PARSER_VERSION = 2

SECTION_MODELS = {
    "experiences": WorkExperience,
    "education": Education,
    "projects": Project,
    "skills": Skill,
    "certifications": Certification,
    "languages": ResumeLanguage,
    "awards": Award,
    "references": Reference,
}

STOP_WORDS = {
    "and",
    "the",
    "with",
    "for",
    "from",
    "that",
    "this",
    "your",
    "you",
    "our",
    "are",
    "will",
    "have",
    "has",
    "job",
    "role",
    "work",
    "years",
    "about",
    "into",
    "their",
    "they",
    "but",
    "not",
    "all",
    "who",
    "aan",
    "als",
    "ben",
    "bij",
    "binnen",
    "dat",
    "de",
    "deze",
    "die",
    "dit",
    "een",
    "en",
    "er",
    "ervaring",
    "het",
    "in",
    "is",
    "je",
    "jij",
    "jou",
    "met",
    "naar",
    "niet",
    "of",
    "om",
    "onze",
    "op",
    "te",
    "van",
    "voor",
    "we",
    "werken",
    "worden",
    "jaar",
    "zijn",
}


def resume_snapshot(resume):
    data = {
        "id": resume.id,
        "title": resume.title,
        "locale": resume.locale,
        "template": resume.template_id,
        "template_settings": resume.template_settings,
        "include_photo": resume.include_photo,
        "personal": model_to_dict(resume.personal) if hasattr(resume, "personal") else {},
        "summary": resume.summary.text if hasattr(resume, "summary") else "",
    }
    data["personal"].pop("id", None)
    data["personal"].pop("resume", None)
    data["personal"].pop("profile_photo", None)
    photo_id = resume.personal.profile_photo_id if hasattr(resume, "personal") else None
    data["personal"]["profile_photo_id"] = str(photo_id) if photo_id else None
    data["personal"]["template_supports_photo"] = bool(resume.template and resume.template.supports_photo)
    for name, model in SECTION_MODELS.items():
        data[name] = []
        for item in model.objects.filter(resume=resume):
            item_data = model_to_dict(item)
            item_data.pop("resume", None)
            data[name].append(item_data)
    return json.loads(json.dumps(data, cls=DjangoJSONEncoder))


def create_version(resume, source="manual", snapshot=None):
    number = (resume.versions.order_by("-version_number").values_list("version_number", flat=True).first() or 0) + 1
    snapshot_data = snapshot if snapshot is not None else resume_snapshot(resume)
    return resume.versions.create(
        version_number=number,
        snapshot=json.loads(json.dumps(snapshot_data, cls=DjangoJSONEncoder)),
        source=source,
    )


def extract_resume_text(upload):
    content = bytes(upload.file_data)
    extension = upload.filename.rsplit(".", 1)[-1].lower()
    if extension == "pdf":
        with pdfplumber.open(io.BytesIO(content)) as document:
            text = "\n".join(page.extract_text(x_tolerance=2, y_tolerance=3) or "" for page in document.pages).strip()
        if text:
            return text
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    if extension == "docx":
        document = DocxDocument(io.BytesIO(content))
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(paragraphs).strip()
    if extension == "txt":
        return content.decode("utf-8", errors="replace").strip()
    raise ValueError("Supported resume formats are PDF, DOCX, and TXT.")


def parse_resume_text(text):
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    email_match = re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", text)
    phone_match = re.search(r"(?:\+?\d[\d\s().-]{7,}\d)", text)
    first_name = _labeled_value(text, ("first name", "voornaam"))
    last_name = _labeled_value(text, ("last name", "surname", "achternaam"))
    if not first_name and not last_name:
        name_line = next((line for line in lines if not _is_resume_heading(line)), "")
        name_parts = name_line.split()
        first_name = name_parts[0] if name_parts else ""
        last_name = " ".join(name_parts[1:]) if len(name_parts) > 1 else ""
    headings = {
        "experience": [
            "experience",
            "employment",
            "work history",
            "work experience",
            "werkervaring",
        ],
        "education": [
            "education",
            "education & training",
            "academic",
            "opleiding",
            "opleidingen",
        ],
        "skills": ["skills", "technologies", "competencies", "vaardigheden", "competenties"],
        "summary": ["summary", "profile", "objective", "profiel", "samenvatting"],
    }
    sections = {key: [] for key in headings}
    current = None
    for line in lines:
        normalized = _normalize_heading(line)
        matched = next(
            (key for key, values in headings.items() if normalized in values),
            None,
        )
        if matched:
            current = matched
        elif current:
            sections[current].append(line)
    skills = []
    for line in sections["skills"]:
        skills.extend(part.strip() for part in re.split(r"[,|•]", line) if part.strip())
    experiences = _parse_experiences(sections["experience"])
    education = _parse_education(sections["education"])
    return {
        "personal": {
            "first_name": first_name,
            "last_name": last_name,
            "email": email_match.group(0) if email_match else "",
            "phone": phone_match.group(0).strip() if phone_match else "",
            "address": _labeled_value(text, ("address", "adres")),
        },
        "summary": " ".join(sections["summary"]),
        "experience_text": "\n".join(sections["experience"]),
        "education_text": "\n".join(sections["education"]),
        "experiences": experiences,
        "education": education,
        "skills": [{"name": skill} for skill in skills[:50]],
        "raw_text": text,
        "parsing_method": "deterministic",
        "parser_version": RESUME_PARSER_VERSION,
    }


def parse_resume_text_with_ai(text, fallback=None):
    fallback = fallback or parse_resume_text(text)
    from apps.ai_services.providers import get_ai_providers

    system_prompt = (
        "You extract structured resume data. Return JSON only. "
        "Never infer, improve, rewrite, or invent facts. Every returned value must be explicitly present "
        "in the source text. Use null, an empty string, or an empty list when unknown."
    )
    user_prompt = json.dumps(
        {
            "task": "Extract this resume into the required schema.",
            "schema": {
                "personal": {
                    "first_name": "",
                    "last_name": "",
                    "email": "",
                    "phone": "",
                    "address": "",
                    "city": "",
                    "country": "",
                    "professional_title": "",
                },
                "summary": "",
                "skills": [{"name": "", "category": ""}],
                "experiences": [
                    {
                        "employer": "",
                        "job_title": "",
                        "start_date": "YYYY-MM-DD or null",
                        "end_date": "YYYY-MM-DD or null",
                        "current": False,
                        "description": "source text only",
                    }
                ],
                "education": [
                    {
                        "institution": "",
                        "degree": "",
                        "field_of_study": "",
                        "start_date": "YYYY-MM-DD or null",
                        "end_date": "YYYY-MM-DD or null",
                        "description": "source text only",
                    }
                ],
                "certifications": [{"name": "", "issuer": ""}],
                "languages": [{"name": "", "proficiency": ""}],
            },
            "resume_text": text,
        },
        ensure_ascii=False,
    )
    for provider_name, provider in get_ai_providers():
        try:
            response = provider.generate(system_prompt, user_prompt)
            parsed = _load_ai_json(response)
            sanitized = _sanitize_ai_resume(parsed, text)
            merged = _merge_parsed_resume(fallback, sanitized)
            merged["parsing_provider"] = provider_name
            return merged
        except Exception as exc:
            logger.warning("Resume parsing with %s failed; trying fallback: %s", provider_name, exc)
    return fallback


def _load_ai_json(value):
    cleaned = value.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.I)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    parsed = json.loads(cleaned)
    if not isinstance(parsed, dict):
        raise ValueError("AI resume parser did not return a JSON object.")
    return parsed


def _sanitize_ai_resume(parsed, source_text):
    personal_fields = {
        "first_name",
        "last_name",
        "email",
        "phone",
        "address",
        "city",
        "country",
        "professional_title",
    }
    personal = {
        key: value.strip()
        for key, value in (parsed.get("personal") or {}).items()
        if key in personal_fields and isinstance(value, str) and _source_contains(source_text, value)
    }
    return {
        "personal": personal,
        "summary": _source_backed_text(parsed.get("summary"), source_text),
        "skills": _sanitize_named_items(parsed.get("skills"), source_text, {"name", "category"}),
        "experiences": _sanitize_resume_items(
            parsed.get("experiences"),
            source_text,
            required=("employer", "job_title"),
            allowed={
                "employer",
                "job_title",
                "start_date",
                "end_date",
                "current",
                "description",
            },
        ),
        "education": _sanitize_resume_items(
            parsed.get("education"),
            source_text,
            required=("institution", "degree"),
            allowed={
                "institution",
                "degree",
                "field_of_study",
                "start_date",
                "end_date",
                "description",
            },
        ),
        "certifications": _sanitize_named_items(parsed.get("certifications"), source_text, {"name", "issuer"}),
        "languages": _sanitize_named_items(parsed.get("languages"), source_text, {"name", "proficiency"}),
        "parsing_method": "ai_assisted",
    }


def _sanitize_named_items(items, source_text, allowed):
    sanitized = []
    for item in items or []:
        if not isinstance(item, dict) or not _source_contains(source_text, item.get("name")):
            continue
        sanitized.append(
            {
                key: value.strip()
                for key, value in item.items()
                if key in allowed and isinstance(value, str) and _source_contains(source_text, value)
            }
        )
    return sanitized


def _sanitize_resume_items(items, source_text, required, allowed):
    sanitized = []
    for item in items or []:
        if not isinstance(item, dict) or not all(_source_contains(source_text, item.get(key)) for key in required):
            continue
        result = {}
        for key, value in item.items():
            if key not in allowed:
                continue
            if key == "current" and isinstance(value, bool):
                result[key] = value
            elif key in {"start_date", "end_date"}:
                result[key] = _source_backed_date(value, source_text)
            elif isinstance(value, str) and _source_contains(source_text, value):
                result[key] = value.strip()
        if all(result.get(key) for key in required):
            sanitized.append(result)
    return sanitized


def _source_backed_text(value, source_text):
    return value.strip() if isinstance(value, str) and _source_contains(source_text, value) else ""


def _source_backed_date(value, source_text):
    if not value:
        return None
    match = re.fullmatch(r"((?:19|20)\d{2})-(\d{2})-(\d{2})", str(value))
    if not match or match.group(1) not in source_text:
        return None
    try:
        return date.fromisoformat(str(value)).isoformat()
    except ValueError:
        return None


def _source_contains(source_text, value):
    if not isinstance(value, str) or not value.strip():
        return False
    normalize = lambda item: re.sub(r"[^\w@.+-]+", " ", item.casefold()).strip()
    return normalize(value) in normalize(source_text)


def _merge_parsed_resume(fallback, ai_data):
    merged = {
        **fallback,
        "parsing_method": ai_data["parsing_method"],
        "parser_version": RESUME_PARSER_VERSION,
    }
    merged["personal"] = {**fallback.get("personal", {}), **ai_data.get("personal", {})}
    for key in ("summary",):
        if ai_data.get(key):
            merged[key] = ai_data[key]
    for key in ("skills", "experiences", "education", "certifications", "languages"):
        if ai_data.get(key):
            merged[key] = ai_data[key]
    return merged


def parsed_resume_is_current(upload):
    return upload.status == "completed" and upload.parsed_json.get("parser_version") == RESUME_PARSER_VERSION


def _normalize_heading(value):
    return re.sub(r"[\s.:]+$", "", value.strip().casefold())


def _is_resume_heading(value):
    return _normalize_heading(value) in {
        "curriculum vitae",
        "resume",
        "cv",
        "personal",
        "personal details",
        "persoonlijke gegevens",
    }


def _labeled_value(text, labels):
    next_label = (
        r"(?=\s+(?:first name|last name|surname|voornaam|achternaam|address|adres|"
        r"phone number|phone|telephone|telefoon|e-?mail|birth date|geboortedatum|"
        r"birth location|gender|marital status)\b|$)"
    )
    for label in labels:
        match = re.search(
            rf"\b{re.escape(label)}\b\s*:?\s*(.+?){next_label}",
            text,
            re.I | re.S,
        )
        if match:
            return re.sub(r"\s+", " ", match.group(1)).strip()
    return ""


def _parse_experiences(lines):
    text = " ".join(lines)
    statements = [
        item.strip()
        for item in re.split(
            r"(?<=[.!?])\s+(?=(?:working|worked|currently working|werkzaam|gewerkt)\b)", text, flags=re.I
        )
        if item.strip()
    ]
    experiences = []
    pattern = re.compile(
        r"(?:currently\s+)?(?:working|worked)\s+at\s+(?P<employer>.+?)\s+as\s+"
        r"(?P<title>.+?)\s+from\s+(?P<start>.+?)(?:\s+to\s+(?P<end>.+?))?[.!]?$",
        re.I,
    )
    for statement in statements:
        match = pattern.search(statement)
        if not match:
            continue
        start_date = _parse_resume_date(match.group("start"))
        end_value = (match.group("end") or "").strip()
        current = not end_value or end_value.casefold() in {"present", "current", "heden", "nu"}
        experiences.append(
            {
                "employer": re.sub(r"\s*\([^)]*\)\s*", " ", match.group("employer")).strip(),
                "job_title": match.group("title").strip(),
                "start_date": start_date.isoformat() if start_date else None,
                "end_date": (None if current else (_parse_resume_date(end_value, end_of_period=True) or None)),
                "current": current,
                "description": statement,
            }
        )
    for item in experiences:
        if isinstance(item["end_date"], date):
            item["end_date"] = item["end_date"].isoformat()
    return experiences


def _parse_education(lines):
    text = " ".join(line for line in lines if not re.match(r"periode\b", line, re.I))
    entries = re.split(r"(?=\b(?:19|20)\d{2}\s*[–—-]\s*(?:19|20)\d{2}\b)", text)
    education = []
    for entry in entries:
        match = re.match(
            r"\s*(?P<start>(?:19|20)\d{2})\s*[–—-]\s*(?P<end>(?:19|20)\d{2})\s+(?P<body>.+)",
            entry,
        )
        if not match:
            continue
        parts = [part.strip() for part in match.group("body").split(",") if part.strip()]
        degree = parts[0] if parts else match.group("body").strip()
        institution = ", ".join(parts[1:]) if len(parts) > 1 else degree
        education.append(
            {
                "institution": institution,
                "degree": degree,
                "start_date": f"{match.group('start')}-01-01",
                "end_date": f"{match.group('end')}-12-31",
            }
        )
    return education


def _parse_resume_date(value, end_of_period=False):
    cleaned = value.strip().rstrip(".")
    year_match = re.search(r"\b((?:19|20)\d{2})\b", cleaned)
    if not year_match:
        return None
    year = int(year_match.group(1))
    months = {
        "jan": 1,
        "january": 1,
        "januari": 1,
        "feb": 2,
        "february": 2,
        "februari": 2,
        "mar": 3,
        "march": 3,
        "maart": 3,
        "apr": 4,
        "april": 4,
        "may": 5,
        "mei": 5,
        "jun": 6,
        "june": 6,
        "juni": 6,
        "jul": 7,
        "july": 7,
        "juli": 7,
        "aug": 8,
        "august": 8,
        "augustus": 8,
        "sep": 9,
        "sept": 9,
        "september": 9,
        "oct": 10,
        "october": 10,
        "okt": 10,
        "oktober": 10,
        "nov": 11,
        "november": 11,
        "dec": 12,
        "december": 12,
    }
    month = next((number for name, number in months.items() if re.search(rf"\b{name}\b", cleaned, re.I)), None)
    if month:
        return date(year, month, monthrange(year, month)[1] if end_of_period else 1)
    return date(year, 12 if end_of_period else 1, 31 if end_of_period else 1)


def parse_upload(upload):
    upload.status = "processing"
    upload.save(update_fields=["status", "updated_at"])
    try:
        text = extract_resume_text(upload)
        parsed = parse_resume_text_with_ai(text)
        upload.extracted_text = text
        upload.parsed_json = parsed
        upload.status = "completed"
        upload.error_message = ""
        upload.save()
        TemporaryResumeUpload.objects.create(
            user=upload.user,
            anonymous_identity=upload.anonymous_identity,
            upload=upload,
            extracted_text=text,
            parsed_json=parsed,
            expires_at=timezone.now() + timedelta(hours=settings.CAREER_SUITE_TEMP_TTL_HOURS),
        )
        return parsed
    except Exception as exc:
        upload.status = "failed"
        upload.error_message = str(exc)
        upload.save(update_fields=["status", "error_message", "updated_at"])
        raise


def keyword_counts(text):
    words = re.findall(r"[A-Za-z][A-Za-z0-9+#.-]{1,}", text.lower())
    return Counter(word for word in words if word not in STOP_WORDS)


def apply_parsed_resume(resume, parsed):
    personal_data = parsed.get("personal", {})
    allowed = {field.name for field in PersonalDetail._meta.fields} - {"id", "resume"}
    PersonalDetail.objects.update_or_create(
        resume=resume, defaults={key: value for key, value in personal_data.items() if key in allowed}
    )
    ResumeSummary.objects.update_or_create(resume=resume, defaults={"text": parsed.get("summary", "")})
    for skill in parsed.get("skills", []):
        if skill.get("name"):
            Skill.objects.get_or_create(resume=resume, name=skill["name"], defaults={"position": 0})
    for position, experience in enumerate(parsed.get("experiences", [])):
        WorkExperience.objects.create(
            resume=resume,
            position=position,
            **{
                key: value
                for key, value in experience.items()
                if key in {"employer", "job_title", "start_date", "end_date", "current", "description"}
            },
        )
    for position, item in enumerate(parsed.get("education", [])):
        Education.objects.create(
            resume=resume,
            position=position,
            **{
                key: value
                for key, value in item.items()
                if key in {"institution", "degree", "field_of_study", "start_date", "end_date", "description"}
            },
        )
    for position, item in enumerate(parsed.get("certifications", [])):
        Certification.objects.create(
            resume=resume,
            position=position,
            **{
                key: value
                for key, value in item.items()
                if key in {"name", "issuer", "issued_at", "expires_at", "credential_id", "credential_url"}
            },
        )
    for position, item in enumerate(parsed.get("languages", [])):
        ResumeLanguage.objects.create(
            resume=resume,
            position=position,
            **{key: value for key, value in item.items() if key in {"name", "proficiency"}},
        )
    create_version(resume, source="upload")
    return resume


def _snapshot_item_key(section, item):
    normalize = lambda value: re.sub(r"\s+", " ", str(value or "").strip()).casefold()
    if section == "skills":
        return (normalize(item.get("name")),)
    if section == "education":
        return (
            normalize(item.get("institution")),
            normalize(item.get("degree")),
            normalize(item.get("field_of_study")),
            str(item.get("start_date") or ""),
        )
    if section == "experiences":
        return (
            normalize(item.get("employer")),
            normalize(item.get("job_title")),
            str(item.get("start_date") or ""),
        )
    if section == "projects":
        return (normalize(item.get("name")), normalize(item.get("role")))
    return None


@transaction.atomic
def create_resume_from_snapshot(
    user,
    source_resume,
    snapshot,
    title_suffix="Optimized",
    version_source="ai_optimization",
):
    resume = source_resume.__class__.objects.create(
        user=source_resume.user,
        anonymous_identity=source_resume.anonymous_identity,
        title=f"{source_resume.title} - {title_suffix}",
        locale=snapshot.get("locale", source_resume.locale),
        template=source_resume.template,
        template_settings=source_resume.template_settings,
        include_photo=source_resume.include_photo,
        source=source_resume.source,
        max_edit_count=source_resume.max_edit_count,
    )
    personal_data = snapshot.get("personal", {}).copy()
    personal_data.pop("profile_photo_id", None)
    personal_data.pop("template_supports_photo", None)
    personal_data.pop("photo_url", None)
    allowed = {field.name for field in PersonalDetail._meta.fields} - {"id", "resume", "profile_photo"}
    personal = PersonalDetail.objects.create(
        resume=resume,
        profile_photo=source_resume.personal.profile_photo if hasattr(source_resume, "personal") else None,
        **{key: value for key, value in personal_data.items() if key in allowed},
    )
    if hasattr(source_resume, "personal"):
        personal.include_photo = source_resume.personal.include_photo
        personal.save(update_fields=["include_photo"])
    ResumeSummary.objects.create(resume=resume, text=snapshot.get("summary", ""))
    for section, model in SECTION_MODELS.items():
        allowed_fields = {field.name for field in model._meta.fields} - {
            "id",
            "resume",
            "created_at",
            "updated_at",
            "normalized_name",
            "duplicate_key",
        }
        seen = set()
        for item in snapshot.get(section, []):
            item_key = _snapshot_item_key(section, item)
            if item_key is not None and item_key in seen:
                continue
            if item_key is not None:
                seen.add(item_key)
            model.objects.create(
                resume=resume,
                **{key: value for key, value in item.items() if key in allowed_fields},
            )
    version = create_version(resume, source=version_source)
    return resume, version
