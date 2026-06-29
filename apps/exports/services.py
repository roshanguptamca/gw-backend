import base64
import io

from django.template.loader import render_to_string
from django.utils import timezone

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from apps.files.models import UserFile
from apps.resumes.services import resume_snapshot

DEFAULT_SECTION_ORDER = [
    "summary",
    "experience",
    "education",
    "skills",
    "projects",
    "certifications",
    "languages",
    "awards",
    "references",
]
FONT_SIZES = {"small": 9, "medium": 10.5, "large": 12}
LINE_SPACING = {"compact": 0.9, "normal": 1.0, "relaxed": 1.2}
EXPORT_LABELS = {
    "en": {
        "summary": "Professional Summary",
        "experience": "Work Experience",
        "education": "Education",
        "skills": "Skills",
        "projects": "Projects",
        "certifications": "Certifications",
        "languages": "Languages",
        "awards": "Awards",
        "references": "References",
        "present": "Present",
    },
    "nl": {
        "summary": "Profiel",
        "experience": "Werkervaring",
        "education": "Opleiding",
        "skills": "Vaardigheden",
        "projects": "Projecten",
        "certifications": "Certificaten",
        "languages": "Talen",
        "awards": "Onderscheidingen",
        "references": "Referenties",
        "present": "Heden",
    },
}


def get_template_settings(resume, overrides=None, template=None):
    template = template or resume.template
    settings = {
        "primary_color": "#1F2937",
        "font_family": "Inter",
        "font_size": "medium",
        "spacing": "normal",
        "include_photo": resume.include_photo,
        "section_order": DEFAULT_SECTION_ORDER,
    }
    if template:
        settings.update(template.default_settings or {})
    settings.update(resume.template_settings or {})
    settings.update(overrides or {})
    settings["include_photo"] = bool(
        settings.get("include_photo", True) and resume.include_photo and template and template.supports_photo
    )
    return settings


def validate_export(resume, output_format):
    if not resume.template_id:
        raise ValueError("Please select a resume template before generating your CV.")
    if output_format not in resume.template.supported_formats:
        raise ValueError(f"The selected template does not support {output_format.upper()} export.")
    snapshot = resume_snapshot(resume)
    personal = snapshot["personal"]
    required = {
        "first_name": personal.get("first_name"),
        "last_name": personal.get("last_name"),
        "email": personal.get("email"),
        "phone": personal.get("phone"),
        "city": personal.get("city"),
        "country": personal.get("country"),
        "professional_title": personal.get("professional_title"),
        "summary": snapshot["summary"],
        "education": snapshot["education"],
        "three_skills": len(snapshot["skills"]) >= 3,
        "language": snapshot["languages"],
        "output_format": output_format in {"pdf", "docx"},
    }
    missing = [key for key, value in required.items() if not value]
    if missing:
        raise ValueError(f"Resume export is missing: {', '.join(missing)}")
    return snapshot


def get_export_photo(resume, settings=None, template=None):
    template = template or resume.template
    resolved = settings or get_template_settings(resume, template=template)
    if (
        not hasattr(resume, "personal")
        or not resume.personal.profile_photo_id
        or not resume.personal.include_photo
        or not resolved.get("include_photo")
        or not template
        or not template.supports_photo
    ):
        return None
    return resume.personal.profile_photo


def render_resume_html(resume, template=None, settings=None, snapshot=None):
    template = template or resume.template
    if not template:
        raise ValueError("Please select a resume template before generating your CV.")
    resolved = get_template_settings(resume, settings, template)
    data = snapshot or resume_snapshot(resume)
    photo = get_export_photo(resume, resolved, template)
    photo_data_uri = None
    if photo:
        encoded = base64.b64encode(bytes(photo.file_data)).decode("ascii")
        photo_data_uri = f"data:{photo.content_type};base64,{encoded}"
    return render_to_string(
        template.html_template,
        {
            "resume": data,
            "resume_template": template,
            "settings": resolved,
            "photo_data_uri": photo_data_uri,
            "labels": EXPORT_LABELS.get(data.get("locale"), EXPORT_LABELS["en"]),
        },
    )


def _add_docx_section(document, data, section):
    labels = EXPORT_LABELS.get(data.get("locale"), EXPORT_LABELS["en"])
    section_map = {
        "experience": (labels["experience"], "experiences"),
        "education": (labels["education"], "education"),
        "skills": (labels["skills"], "skills"),
        "projects": (labels["projects"], "projects"),
        "certifications": (labels["certifications"], "certifications"),
        "languages": (labels["languages"], "languages"),
        "awards": (labels["awards"], "awards"),
        "references": (labels["references"], "references"),
    }
    if section == "summary":
        if data["summary"]:
            document.add_heading(labels["summary"], level=1)
            document.add_paragraph(data["summary"])
        return
    heading, key = section_map[section]
    if not data.get(key):
        return
    document.add_heading(heading, level=1)
    if section == "skills":
        document.add_paragraph(", ".join(item["name"] for item in data[key]))
        return
    for item in data[key]:
        title = item.get("job_title") or item.get("degree") or item.get("name") or item.get("title") or heading
        document.add_heading(str(title), level=2)
        detail = (
            item.get("employer")
            or item.get("institution")
            or item.get("issuer")
            or item.get("company")
            or item.get("proficiency")
        )
        if detail:
            document.add_paragraph(str(detail))
        if item.get("description"):
            document.add_paragraph(item["description"])


def generate_docx(resume):
    data = validate_export(resume, "docx")
    settings = get_template_settings(resume)
    document = Document()
    personal = data["personal"]
    photo = get_export_photo(resume, settings=settings)
    if photo:
        document.add_picture(io.BytesIO(bytes(photo.file_data)), width=Inches(1.2))
    document.add_heading(f"{personal['first_name']} {personal['last_name']}", 0)
    document.add_paragraph(personal["professional_title"])
    document.add_paragraph(" | ".join([personal["email"], personal["phone"], personal["city"], personal["country"]]))
    for section in settings["section_order"]:
        _add_docx_section(document, data, section)
    color = RGBColor.from_string(settings["primary_color"].lstrip("#"))
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2"):
        style = document.styles[style_name]
        style.font.name = settings["font_family"]
        style.font.size = Pt(FONT_SIZES[settings["font_size"]])
        if style_name != "Normal":
            style.font.color.rgb = color
    for paragraph in document.paragraphs:
        paragraph.paragraph_format.line_spacing = LINE_SPACING[settings["spacing"]]
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def generate_pdf(resume):
    data = validate_export(resume, "pdf")
    from weasyprint import HTML

    return HTML(string=render_resume_html(resume, snapshot=data)).write_pdf()


def save_export(resume, output_format):
    if output_format not in {"pdf", "docx"}:
        raise ValueError("Supported export formats are PDF and DOCX.")
    content = generate_pdf(resume) if output_format == "pdf" else generate_docx(resume)
    filename = f"{resume.title.replace(' ', '_')}_{timezone.now():%Y%m%d}.{output_format}"
    return UserFile.objects.create(
        user=resume.user,
        anonymous_identity=resume.anonymous_identity,
        filename=filename,
        content_type=(
            "application/pdf"
            if output_format == "pdf"
            else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        file_size=len(content),
        file_data=content,
        purpose="resume_export",
    )
