import logging

import pandas as pd
import pytesseract
from docx import Document as DocxDocument
from PIL import Image
from pypdf import PdfReader

logger = logging.getLogger(__name__)


def extract_pdf(path):
    """Extract text from PDF file."""
    try:
        reader = PdfReader(path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text.strip()
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise ValueError(f"Failed to extract PDF: {str(e)}")


def extract_docx(path):
    """Extract text from DOCX file. Also attempted for legacy .doc files."""
    try:
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs)
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Failed to extract DOCX: {str(e)}")


def extract_doc(path):
    """Extract text from legacy .doc (Word 97-2003) files."""
    # First try python-docx — some .doc files are actually Office Open XML
    try:
        return extract_docx(path)
    except Exception:
        pass

    # Try antiword (CLI tool, may not be installed)
    try:
        import subprocess

        result = subprocess.run(["antiword", path], capture_output=True, text=True, timeout=30)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except (FileNotFoundError, subprocess.TimeoutExpired, Exception):
        pass

    # Try textract if available
    try:
        import textract

        text = textract.process(path).decode("utf-8", errors="ignore")
        if text.strip():
            return text.strip()
    except (ImportError, Exception):
        pass

    raise ValueError(
        "Cannot extract text from this .doc file. "
        "Please convert it to .docx (Save As → Word Document) and re-upload."
    )


def extract_image(path):
    """Extract text from image using OCR."""
    try:
        text = pytesseract.image_to_string(Image.open(path))
        return text.strip()
    except Exception as e:
        logger.error(f"Image extraction failed: {e}")
        raise ValueError(f"Failed to extract image text: {str(e)}")


def extract_txt(path):
    """Extract text from plain text file."""
    try:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
        return text.strip()
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(path, "r", encoding="latin-1") as f:
                text = f.read()
            return text.strip()
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            raise ValueError(f"Failed to extract text file: {str(e)}")
    except Exception as e:
        logger.error(f"TXT extraction failed: {e}")
        raise ValueError(f"Failed to extract text file: {str(e)}")


def extract_csv(path, max_rows=1000):
    """
    Extract text from CSV file.
    Converts CSV to readable text format with column headers.

    Args:
        path: Path to CSV file
        max_rows: Maximum number of rows to process (default: 1000)

    Returns:
        Formatted text representation of CSV data
    """
    try:
        # Read CSV with pandas
        df = pd.read_csv(path, nrows=max_rows)

        # Check if empty
        if df.empty:
            return "Empty CSV file"

        # Build text representation
        text_parts = []

        # Add column info
        text_parts.append(f"CSV File with {len(df.columns)} columns and {len(df)} rows\n")
        text_parts.append(f"Columns: {', '.join(df.columns.tolist())}\n")

        # Add data summary
        text_parts.append("\nData Summary:")
        text_parts.append(df.describe(include="all").to_string())

        # Add first few rows
        text_parts.append("\n\nSample Data (first 10 rows):")
        text_parts.append(df.head(10).to_string(index=False))

        return "\n".join(text_parts)

    except pd.errors.EmptyDataError:
        return "Empty CSV file"
    except Exception as e:
        logger.error(f"CSV extraction failed: {e}")
        raise ValueError(f"Failed to extract CSV: {str(e)}")


def extract_xlsx(path, max_rows=1000):
    """
    Extract text from Excel (XLSX) file.
    Processes all sheets and converts to readable text format.

    Args:
        path: Path to XLSX file
        max_rows: Maximum number of rows per sheet (default: 1000)

    Returns:
        Formatted text representation of Excel data
    """
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(path)
        text_parts = []

        text_parts.append(f"Excel File with {len(excel_file.sheet_names)} sheet(s)\n")
        text_parts.append(f"Sheets: {', '.join(excel_file.sheet_names)}\n")

        # Process each sheet
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(path, sheet_name=sheet_name, nrows=max_rows)

            if df.empty:
                text_parts.append(f"\n--- Sheet: {sheet_name} (Empty) ---")
                continue

            text_parts.append(f"\n--- Sheet: {sheet_name} ---")
            text_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
            text_parts.append(f"Rows: {len(df)}")

            # Add data summary
            text_parts.append("\nData Summary:")
            text_parts.append(df.describe(include="all").to_string())

            # Add first few rows
            text_parts.append("\nSample Data (first 5 rows):")
            text_parts.append(df.head(5).to_string(index=False))

        return "\n".join(text_parts)

    except Exception as e:
        logger.error(f"XLSX extraction failed: {e}")
        raise ValueError(f"Failed to extract Excel file: {str(e)}")


# File type mapping
EXTRACTORS = {
    "pdf": extract_pdf,
    "docx": extract_docx,
    "doc": extract_doc,
    "txt": extract_txt,
    "csv": extract_csv,
    "xlsx": extract_xlsx,
    "xls": extract_xlsx,
    "png": extract_image,
    "jpg": extract_image,
    "jpeg": extract_image,
}


def extract_text(file_path, file_type=None):
    """
    Universal text extractor that automatically detects file type.

    Args:
        file_path: Path to the file
        file_type: Optional file type hint (e.g., 'pdf', 'docx')

    Returns:
        Extracted text content
    """
    if not file_type:
        # Try to detect from file extension
        import os

        _, ext = os.path.splitext(file_path)
        file_type = ext.lower().replace(".", "")

    extractor = EXTRACTORS.get(file_type.lower())
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")

    return extractor(file_path)
