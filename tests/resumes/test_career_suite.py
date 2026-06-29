import io
import zipfile
from unittest.mock import Mock, patch

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import override_settings

import pytest
import requests
from docx import Document
from PIL import Image
from rest_framework.test import APIClient

from apps.exports.services import generate_docx, get_export_photo, render_resume_html
from apps.files.models import UserFile
from apps.jobs.models import JobDescription, TemporaryJobDescription
from apps.jobs.services import analyze_match, parse_job_text, parse_job_url
from apps.resumes.models import (
    AnonymousResumeIdentity,
    Education,
    OptimizedResume,
    PersonalDetail,
    Project,
    Resume,
    ResumeSummary,
    ResumeUpload,
    Skill,
    TemporaryGeneratedResume,
    WorkExperience,
)
from apps.resumes.services import create_resume_from_snapshot, parse_resume_text, parse_resume_text_with_ai
from apps.templates_app.models import ResumeTemplate


@pytest.fixture
def user(db):
    return get_user_model().objects.create_user("career-user", "career@example.com", "password")


@pytest.fixture
def other_user(db):
    return get_user_model().objects.create_user("other-user", "other@example.com", "password")


@pytest.fixture
def client(user):
    api_client = APIClient()
    api_client.force_authenticate(user)
    return api_client


@pytest.mark.django_db
def test_resume_crud_is_owner_scoped(client, user, other_user):
    create_response = client.post("/api/resumes/", {"title": "Backend Engineer", "locale": "en"}, format="json")
    assert create_response.status_code == 201
    resume_id = create_response.data["id"]

    other_resume = Resume.objects.create(user=other_user, title="Private")
    list_response = client.get("/api/resumes/")

    assert list_response.status_code == 200
    assert [item["id"] for item in list_response.data] == [resume_id]
    assert client.get(f"/api/resumes/{other_resume.id}/").status_code == 404


@pytest.mark.django_db
def test_anonymous_resume_limit_edit_limit_and_owner_scope():
    anonymous = APIClient()
    owner = {"owner_email": "anon@example.com", "owner_phone": "+31 6 12345678"}
    created = anonymous.post(
        "/api/resumes/",
        {"title": "Anonymous Resume", **owner},
        format="json",
        REMOTE_ADDR="203.0.113.10",
    )
    assert created.status_code == 201
    resume_id = created.data["id"]
    assert created.data["usage"]["remaining_edits"] == 10

    duplicate = anonymous.post(
        "/api/resumes/",
        {"title": "Second Resume", **owner},
        format="json",
        REMOTE_ADDR="203.0.113.10",
    )
    assert duplicate.status_code == 400
    assert "only 1 resume" in str(duplicate.data)

    for edit_number in range(1, 11):
        response = anonymous.put(
            f"/api/resumes/{resume_id}/",
            {"title": f"Anonymous Resume {edit_number}", "locale": "en"},
            format="json",
            REMOTE_ADDR="203.0.113.10",
        )
        assert response.status_code == 200
        assert response.data["usage"]["edit_count"] == edit_number

    blocked = anonymous.put(
        f"/api/resumes/{resume_id}/",
        {"title": "Eleventh edit", "locale": "en"},
        format="json",
        REMOTE_ADDR="203.0.113.10",
    )
    assert blocked.status_code == 400
    assert "edit limit reached" in str(blocked.data)

    other = APIClient()
    assert other.get(f"/api/resumes/{resume_id}/", REMOTE_ADDR="203.0.113.11").status_code == 404
    assert anonymous.get(f"/api/resumes/{resume_id}/", REMOTE_ADDR="203.0.113.10").status_code == 200


@pytest.mark.django_db
def test_anonymous_preview_ownership_lookup_does_not_write_identity():
    browser = APIClient()
    created = browser.post(
        "/api/resumes/",
        {
            "title": "Anonymous Preview",
            "owner_email": "preview@example.com",
            "owner_phone": "+31612345678",
        },
        format="json",
        REMOTE_ADDR="203.0.113.20",
    )
    resume = Resume.objects.get(id=created.data["id"])
    template = ResumeTemplate.objects.create(slug="preview-safe", name="Preview Safe")
    before = resume.anonymous_identity.last_seen_at

    with patch("apps.resumes.models.AnonymousResumeIdentity.save") as identity_save:
        response = browser.post(
            f"/api/resumes/{resume.id}/preview/",
            {"template_id": template.id},
            format="json",
            REMOTE_ADDR="203.0.113.20",
        )

    assert response.status_code == 200
    identity_save.assert_not_called()
    resume.anonymous_identity.refresh_from_db()
    assert resume.anonymous_identity.last_seen_at == before


@pytest.mark.django_db
def test_registered_resume_limit_and_anonymous_claim(user):
    registered = APIClient()
    registered.force_authenticate(user)
    for number in range(3):
        assert registered.post("/api/resumes/", {"title": f"Resume {number}"}, format="json").status_code == 201
    fourth = registered.post("/api/resumes/", {"title": "Resume 4"}, format="json")
    assert fourth.status_code == 400
    assert "up to 3 resumes" in str(fourth.data)

    claimant = get_user_model().objects.create_user("claimant", "claim@example.com", "password")
    browser = APIClient()
    anonymous_resume = browser.post(
        "/api/resumes/",
        {
            "title": "Claim me",
            "owner_email": "claim@example.com",
            "owner_phone": "+31687654321",
        },
        format="json",
        REMOTE_ADDR="198.51.100.8",
    )
    assert anonymous_resume.status_code == 201
    identity_id = Resume.objects.get(id=anonymous_resume.data["id"]).anonymous_identity_id
    assert AnonymousResumeIdentity.objects.filter(id=identity_id).exists()

    browser.force_authenticate(claimant)
    claimed = browser.post(
        "/api/resumes/claim-anonymous/",
        {"email": "claim@example.com"},
        format="json",
        REMOTE_ADDR="198.51.100.8",
    )
    assert claimed.status_code == 200
    resume = Resume.objects.get(id=anonymous_resume.data["id"])
    assert resume.user == claimant
    assert resume.anonymous_identity is None
    assert resume.is_claimed is True
    assert resume.source == "registered"


@pytest.mark.django_db
def test_resume_sections_and_job_match(client):
    resume = client.post("/api/resumes/", {"title": "Engineer"}, format="json").data
    resume_id = resume["id"]
    personal = {
        "first_name": "Ada",
        "last_name": "Lovelace",
        "email": "ada@example.com",
        "phone": "+31 600000000",
        "city": "Amsterdam",
        "country": "Netherlands",
        "professional_title": "Python Engineer",
    }
    assert client.put(f"/api/resumes/{resume_id}/personal/", personal, format="json").status_code == 200
    assert (
        client.put(
            f"/api/resumes/{resume_id}/summary/", {"text": "Python backend engineer."}, format="json"
        ).status_code
        == 200
    )
    for skill in ("Python", "Django", "PostgreSQL"):
        assert (
            client.post(
                f"/api/resumes/{resume_id}/skills/",
                {"skill_name": skill, "category": "Technical"},
                format="json",
            ).status_code
            == 201
        )
    assert (
        client.post(
            f"/api/resumes/{resume_id}/education/",
            {"institution": "University", "degree": "BSc Computer Science"},
            format="json",
        ).status_code
        == 201
    )
    assert (
        client.post(
            f"/api/resumes/{resume_id}/languages/",
            {"name": "English", "proficiency": "Fluent"},
            format="json",
        ).status_code
        == 201
    )

    parsed_job = client.post(
        "/api/jobs/parse-text/",
        {"title": "Python Engineer", "text": "Job title: Python Engineer. Python Django PostgreSQL API"},
        format="json",
    )
    assert parsed_job.status_code == 201
    job_id = parsed_job.data["job_description"]["id"]
    match = client.post(
        "/api/job-match/analyze/",
        {"resume_id": resume_id, "job_description_id": job_id},
        format="json",
    )
    assert match.status_code == 201
    assert float(match.data["overall_score"]) > 0
    assert "ats_report" in match.data
    optimized = client.post(
        f"/api/job-match/{match.data['id']}/optimize/",
        {"target_score": 90, "confirmed_skills": [], "declined_skills": []},
        format="json",
    )
    assert optimized.status_code == 201
    assert OptimizedResume.objects.filter(job_match_id=match.data["id"]).exists()
    assert (
        Resume.objects.get(id=optimized.data["optimized_resume_id"]).versions.filter(source="ai_optimization").exists()
    )
    assert TemporaryGeneratedResume.objects.filter(source_resume_id=resume_id).count() == 2


@pytest.mark.django_db
def test_docx_upload_parse_and_create_resume(client):
    document = Document()
    document.add_paragraph("Grace Hopper")
    document.add_paragraph("grace@example.com")
    document.add_heading("Summary")
    document.add_paragraph("Compiler engineer and computer scientist.")
    document.add_heading("Skills")
    document.add_paragraph("Python, COBOL, Leadership")
    content = io.BytesIO()
    document.save(content)
    upload = SimpleUploadedFile(
        "resume.docx",
        content.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    upload_response = client.post("/api/resumes/upload/", {"file": upload}, format="multipart")
    assert upload_response.status_code == 201
    assert upload_response.data["status"] == "completed"

    parse_response = client.post(
        "/api/resumes/parse/",
        {"upload_id": upload_response.data["id"], "create_resume": True},
        format="json",
    )
    assert parse_response.status_code == 200
    assert parse_response.data["resume"]["personal"]["email"] == "grace@example.com"
    assert len(parse_response.data["resume"]["skills"]) == 3


@pytest.mark.django_db
def test_pdf_layout_resume_parser_persists_personal_education_and_experience(client):
    extracted_text = """Curriculum Vitae
Personal
First name Rati
Last name Gupta
Address: Vuurdoornpark 2
2724HE, Zoetermeer
Phone number 0647696248
E-mail guptarati024@gmail.com
Birth date 23-05-1986
Education & Training
Periode Naam opleiding en school
2008 – 2011 Bachelor of science, Madhya Pradesh Bhoj University,
Bhopal
2001 – 2008 Secondary school – High school (Middlebare school)
Work experience.
Working at Casa (https://casaschool.nl/) as Kitchen assistant from Sep-2024.
Worked at Disha computer Vashi Navi Mumbai India as marketing adviser from June-2013 to Dec-
2015"""
    parsed = parse_resume_text(extracted_text)

    assert parsed["personal"] == {
        "first_name": "Rati",
        "last_name": "Gupta",
        "email": "guptarati024@gmail.com",
        "phone": "0647696248",
        "address": "Vuurdoornpark 2 2724HE, Zoetermeer",
    }
    assert parsed["education"][0]["degree"] == "Bachelor of science"
    assert parsed["education"][0]["institution"] == "Madhya Pradesh Bhoj University, Bhopal"
    assert parsed["experiences"][0]["employer"] == "Casa"
    assert parsed["experiences"][0]["job_title"] == "Kitchen assistant"
    assert parsed["experiences"][0]["current"] is True
    assert parsed["experiences"][1]["end_date"] == "2015-12-31"

    upload = ResumeUpload.objects.create(
        user=client.handler._force_user,
        filename="rati_cv.pdf",
        content_type="application/pdf",
        file_size=100,
        file_data=b"already parsed",
        status="completed",
        extracted_text=extracted_text,
        parsed_json=parsed,
    )
    response = client.post(
        "/api/resumes/parse/",
        {"upload_id": str(upload.id), "create_resume": True},
        format="json",
    )

    assert response.status_code == 200
    resume = Resume.objects.get(id=response.data["resume"]["id"])
    assert resume.personal.first_name == "Rati"
    assert Education.objects.filter(resume=resume).count() == 2
    assert WorkExperience.objects.filter(resume=resume).count() == 2


@pytest.mark.django_db
def test_parse_endpoint_reprocesses_completed_upload_from_old_parser(client):
    document = Document()
    document.add_paragraph("Grace Hopper")
    document.add_paragraph("grace@example.com")
    document.add_heading("Skills")
    document.add_paragraph("COBOL")
    content = io.BytesIO()
    document.save(content)
    upload = ResumeUpload.objects.create(
        user=client.handler._force_user,
        filename="legacy.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        file_size=len(content.getvalue()),
        file_data=content.getvalue(),
        status="completed",
        parsed_json={"personal": {"first_name": "Wrong"}},
    )

    response = client.post("/api/resumes/parse/", {"upload_id": str(upload.id)}, format="json")

    assert response.status_code == 200
    assert response.data["parsed_json"]["personal"]["first_name"] == "Grace"
    assert response.data["parsed_json"]["parser_version"] == 2


@override_settings(AI_PROVIDER_FALLBACKS="gemini,openai")
def test_ai_resume_parser_structures_source_facts_and_rejects_inventions():
    source = (
        "Rati Gupta can be reached at guptarati024@gmail.com. "
        "She worked at Casa as Kitchen assistant from September 2024. "
        "Her listed skill is Customer Service."
    )
    provider = Mock()
    provider.generate.return_value = """{
      "personal": {
        "first_name": "Rati",
        "last_name": "Gupta",
        "email": "guptarati024@gmail.com"
      },
      "summary": "",
      "skills": [
        {"name": "Customer Service", "category": ""},
        {"name": "AWS", "category": "Technical"}
      ],
      "experiences": [{
        "employer": "Casa",
        "job_title": "Kitchen assistant",
        "start_date": "2024-09-01",
        "end_date": null,
        "current": true,
        "description": "worked at Casa as Kitchen assistant from September 2024"
      }],
      "education": [],
      "certifications": [],
      "languages": []
    }"""

    with patch("apps.ai_services.providers.get_ai_providers", return_value=[("gemini", provider)]):
        parsed = parse_resume_text_with_ai(source)

    assert parsed["parsing_method"] == "ai_assisted"
    assert parsed["parsing_provider"] == "gemini"
    assert parsed["personal"]["first_name"] == "Rati"
    assert parsed["experiences"][0]["employer"] == "Casa"
    assert parsed["experiences"][0]["start_date"] == "2024-09-01"
    assert [item["name"] for item in parsed["skills"]] == ["Customer Service"]


@override_settings(AI_PROVIDER_FALLBACKS="gemini,openai")
def test_ai_resume_parser_uses_openai_when_gemini_fails():
    gemini = Mock()
    gemini.generate.side_effect = RuntimeError("quota exceeded")
    openai = Mock()
    openai.generate.return_value = """{
      "personal": {"first_name": "Grace", "last_name": "Hopper", "email": "grace@example.com"},
      "summary": "",
      "skills": [{"name": "COBOL", "category": ""}],
      "experiences": [],
      "education": [],
      "certifications": [],
      "languages": []
    }"""
    source = "Grace Hopper\ngrace@example.com\nSkills\nCOBOL"

    with patch(
        "apps.ai_services.providers.get_ai_providers",
        return_value=[("gemini", gemini), ("openai", openai)],
    ):
        parsed = parse_resume_text_with_ai(source)

    assert parsed["parsing_method"] == "ai_assisted"
    assert parsed["parsing_provider"] == "openai"
    assert parsed["personal"]["email"] == "grace@example.com"
    assert parsed["skills"] == [{"name": "COBOL"}]


@override_settings(AI_PROVIDER_FALLBACKS="gemini,openai")
def test_ai_resume_parser_uses_deterministic_fallback_when_all_providers_fail():
    gemini = Mock()
    gemini.generate.side_effect = RuntimeError("quota exceeded")
    openai = Mock()
    openai.generate.side_effect = RuntimeError("provider unavailable")
    source = "Grace Hopper\ngrace@example.com\nSkills\nCOBOL"

    with patch(
        "apps.ai_services.providers.get_ai_providers",
        return_value=[("gemini", gemini), ("openai", openai)],
    ):
        parsed = parse_resume_text_with_ai(source)

    assert parsed["parsing_method"] == "deterministic"
    assert "parsing_provider" not in parsed
    assert parsed["personal"]["email"] == "grace@example.com"


@pytest.mark.django_db
def test_generate_professional_summary_uses_resume_facts_and_target_title(client, user):
    resume = Resume.objects.create(user=user, title="Support Resume", locale="en")
    PersonalDetail.objects.create(resume=resume, professional_title="Kitchen Assistant")
    WorkExperience.objects.create(
        resume=resume,
        employer="Casa",
        job_title="Kitchen Assistant",
        start_date="2024-09-01",
        description="Supported daily kitchen operations.",
    )
    Skill.objects.create(resume=resume, name="Customer Service", category="Functional")
    provider = Mock()
    provider.generate.return_value = (
        "Customer-focused professional with experience supporting daily kitchen operations at Casa. "
        "Brings customer service skills to a Customer Support Specialist role."
    )

    with patch("apps.resumes.summary_generator.get_ai_providers", return_value=[("gemini", provider)]):
        response = client.post(
            f"/api/resumes/{resume.id}/generate-summary/",
            {"job_title": "Customer Support Specialist", "language": "en"},
            format="json",
        )

    assert response.status_code == 200
    assert response.data["provider"] == "gemini"
    assert response.data["requires_user_review"] is True
    assert "Customer Support Specialist" in response.data["summary"]
    prompt = provider.generate.call_args.args[1]
    assert '"target_job_title": "Customer Support Specialist"' in prompt
    assert '"employer": "Casa"' in prompt
    assert "years of experience" not in response.data["summary"]
    assert not ResumeSummary.objects.filter(resume=resume).exists()


@pytest.mark.django_db
def test_generate_professional_summary_falls_back_in_dutch(client, user):
    resume = Resume.objects.create(user=user, title="Support Resume", locale="nl")
    PersonalDetail.objects.create(resume=resume, professional_title="Keukenassistent")
    Skill.objects.create(resume=resume, name="Klantenservice", category="Functioneel")
    gemini = Mock()
    gemini.generate.side_effect = RuntimeError("quota exceeded")
    openai = Mock()
    openai.generate.side_effect = RuntimeError("provider unavailable")

    with patch(
        "apps.resumes.summary_generator.get_ai_providers",
        return_value=[("gemini", gemini), ("openai", openai)],
    ):
        response = client.post(
            f"/api/resumes/{resume.id}/generate-summary/",
            {"job_title": "Klantenservice Medewerker", "language": "nl"},
            format="json",
        )

    assert response.status_code == 200
    assert response.data["provider"] == "deterministic"
    assert response.data["generated_by_ai"] is False
    assert "Klantenservice Medewerker" in response.data["summary"]
    assert "Klantenservice" in response.data["summary"]


@pytest.mark.django_db
def test_generate_skill_suggestions_requires_user_confirmation_and_excludes_saved_skills(client, user):
    resume = Resume.objects.create(user=user, title="Support Resume", locale="en")
    Skill.objects.create(resume=resume, name="Communication", category="Soft Skill")
    provider = Mock()
    provider.generate.return_value = """{
      "skills": [
        {"name": "Communication", "category": "Soft Skill"},
        {"name": "Customer Service", "category": "Functional"},
        {"name": "CRM", "category": "Tool"}
      ]
    }"""

    with patch("apps.resumes.summary_generator.get_ai_providers", return_value=[("gemini", provider)]):
        response = client.post(
            f"/api/resumes/{resume.id}/generate-skills/",
            {"job_title": "Customer Support Specialist", "language": "en"},
            format="json",
        )

    assert response.status_code == 200
    assert response.data["provider"] == "gemini"
    assert response.data["requires_user_confirmation"] is True
    assert [item["name"] for item in response.data["skills"]] == ["Customer Service", "CRM"]
    assert all(item["status"] == "needs_confirmation" for item in response.data["skills"])
    assert list(resume.skill_set.values_list("name", flat=True)) == ["Communication"]


@pytest.mark.django_db
def test_generate_skill_suggestions_has_dutch_fallback(client, user):
    resume = Resume.objects.create(user=user, title="Support Resume", locale="nl")
    gemini = Mock()
    gemini.generate.side_effect = RuntimeError("quota exceeded")
    openai = Mock()
    openai.generate.side_effect = RuntimeError("provider unavailable")

    with patch(
        "apps.resumes.summary_generator.get_ai_providers",
        return_value=[("gemini", gemini), ("openai", openai)],
    ):
        response = client.post(
            f"/api/resumes/{resume.id}/generate-skills/",
            {"job_title": "Customer Support Specialist", "language": "nl"},
            format="json",
        )

    assert response.status_code == 200
    assert response.data["provider"] == "deterministic"
    assert response.data["skills"][0]["name"] == "Klantenservice"
    assert not resume.skill_set.exists()


@pytest.mark.django_db
def test_docx_export_and_protected_download(client, user, other_user):
    template, _ = ResumeTemplate.objects.get_or_create(slug="classic", defaults={"name": "Classic"})
    resume = Resume.objects.create(user=user, title="Complete Resume", template=template)
    client.put(
        f"/api/resumes/{resume.id}/personal/",
        {
            "first_name": "Ada",
            "last_name": "Lovelace",
            "email": "ada@example.com",
            "phone": "123456789",
            "city": "London",
            "country": "UK",
            "professional_title": "Engineer",
        },
        format="json",
    )
    client.put(f"/api/resumes/{resume.id}/summary/", {"text": "Engineering leader."}, format="json")
    client.post(
        f"/api/resumes/{resume.id}/education/",
        {"institution": "University", "degree": "Computer Science"},
        format="json",
    )
    for name in ("Python", "Django", "SQL"):
        client.post(
            f"/api/resumes/{resume.id}/skills/",
            {"skill_name": name, "category": "Technical"},
            format="json",
        )
    client.post(
        f"/api/resumes/{resume.id}/languages/",
        {"name": "English", "proficiency": "Fluent"},
        format="json",
    )

    export_response = client.post(f"/api/resumes/{resume.id}/export/docx/", {}, format="json")
    assert export_response.status_code == 201
    file_id = export_response.data["file_id"]
    assert UserFile.objects.filter(id=file_id, user=user).exists()
    assert client.get(f"/api/files/download/{file_id}/").status_code == 200

    pdf_response = client.post(f"/api/resumes/{resume.id}/export/pdf/", {}, format="json")
    assert pdf_response.status_code == 201
    pdf_download = client.get(f"/api/files/download/{pdf_response.data['file_id']}/")
    assert pdf_download.status_code == 200
    assert pdf_download.content.startswith(b"%PDF")

    other_client = APIClient()
    other_client.force_authenticate(other_user)
    assert other_client.get(f"/api/files/download/{file_id}/").status_code == 404


def test_job_url_redirect_is_validated_before_following():
    redirect = Mock()
    redirect.is_redirect = True
    redirect.is_permanent_redirect = False
    redirect.headers = {"Location": "http://127.0.0.1/admin"}

    def validate(parsed_url):
        if parsed_url.hostname == "127.0.0.1":
            raise ValueError("Private network address")

    with (
        patch("apps.jobs.services._validate_public_url", side_effect=validate),
        patch("apps.jobs.services.requests.get", return_value=redirect) as get,
    ):
        with pytest.raises(ValueError, match="Private network"):
            parse_job_url("https://jobs.example.com/role")

    assert get.call_count == 1


def test_linkedin_job_url_extracts_metadata_and_real_skills():
    html = """
        <html>
          <head>
            <meta property="og:title"
                  content="Alliander hiring Technical Solution Lead in Arnhem, Gelderland, Netherlands | LinkedIn">
          </head>
          <body>
            De energietransitie is onmogelijk zonder jou. Je hebt ervaring met AWS, Kafka,
            NestJS, Angular, Kubernetes, TypeScript, SAP S4/HANA en Snowflake.
          </body>
        </html>
    """
    response = Mock()
    response.is_redirect = False
    response.is_permanent_redirect = False
    response.content = html.encode()
    response.text = html
    response.raise_for_status.return_value = None

    with (
        patch("apps.jobs.services._validate_public_url"),
        patch("apps.jobs.services.requests.get", return_value=response) as get,
    ):
        parsed = parse_job_url("https://www.linkedin.com/jobs/view/123")

    assert get.call_count == 1
    assert parsed["title"] == "Technical Solution Lead"
    assert parsed["company"] == "Alliander"
    assert parsed["location"] == "Arnhem, Gelderland, Netherlands"
    assert {"AWS", "Kafka", "NestJS", "Angular", "Kubernetes", "TypeScript", "SAP", "Snowflake"} <= set(
        parsed["required_skills"]
    )
    assert "en" not in parsed["required_skills"]
    assert "van" not in parsed["keywords"]


@pytest.mark.django_db
def test_english_and_dutch_export_headings(client, user):
    template = seed_template_catalog()[0]
    resume = Resume.objects.create(user=user, title="Bilingual", template=template, locale="en")
    complete_resume(client, resume)
    WorkExperience.objects.create(resume=resume, employer="GuideWisey", job_title="Engineer", start_date="2024-01-01")
    assert "Work Experience" in render_resume_html(resume)
    assert "Professional Summary" in render_resume_html(resume)

    resume.locale = "nl"
    resume.save(update_fields=["locale"])
    html = render_resume_html(resume)
    assert "Werkervaring" in html
    assert "Profiel" in html
    document = Document(io.BytesIO(generate_docx(resume)))
    assert "Opleiding" in "\n".join(paragraph.text for paragraph in document.paragraphs)


@pytest.mark.django_db
def test_dutch_job_skill_matches_english_resume_skill(user):
    resume = Resume.objects.create(user=user, title="Cross-language", locale="en")
    Skill.objects.create(resume=resume, name="Communication", category="Soft Skill")
    parsed = parse_job_text("Vereiste vaardigheden: Communicatie, Leiderschap. Verantwoordelijkheden: samenwerken.")
    job = JobDescription.objects.create(user=user, raw_text=parsed["raw_text"], parsed_json=parsed, language="nl")
    match = analyze_match(resume, job, user, report_language="nl")
    assert "communication" in match.result_json["matched_skills"]
    assert match.report_language == "nl"


@pytest.mark.django_db
def test_invalid_resume_and_optimizer_languages_are_rejected(client, user):
    resume = Resume.objects.create(user=user, title="Language validation")
    invalid_resume = client.put(
        f"/api/resumes/{resume.id}/",
        {"title": resume.title, "locale": "de"},
        format="json",
    )
    assert invalid_resume.status_code == 400

    job = JobDescription.objects.create(user=user, raw_text="Python", parsed_json={"required_skills": ["Python"]})
    match = analyze_match(resume, job, user)
    invalid_optimize = client.post(
        f"/api/job-match/{match.id}/optimize/",
        {"target_score": 80, "output_language": "de"},
        format="json",
    )
    assert invalid_optimize.status_code == 400
    assert (
        client.post(
            "/api/jobs/parse-text/",
            {"text": "Python", "language": "de"},
            format="json",
        ).status_code
        == 400
    )


@pytest.mark.django_db
def test_optimizer_persists_dutch_output_language(client, user):
    resume = Resume.objects.create(user=user, title="Nederlandse optimalisatie", locale="en")
    PersonalDetail.objects.create(resume=resume, professional_title="Engineer")
    Skill.objects.create(resume=resume, name="Communication", category="Soft Skill")
    job = JobDescription.objects.create(
        user=user,
        raw_text="Vereiste vaardigheden: Communicatie en Leiderschap",
        parsed_json={"required_skills": ["communication", "leadership"], "keywords": ["communicatie", "leiderschap"]},
        language="nl",
    )
    match = analyze_match(resume, job, user, report_language="nl")
    response = client.post(
        f"/api/job-match/{match.id}/optimize/",
        {"target_score": 80, "output_language": "nl", "confirmed_skills": [], "declined_skills": ["leadership"]},
        format="json",
    )
    assert response.status_code == 201
    optimized = OptimizedResume.objects.get(id=response.data["id"])
    optimized.optimized_resume.refresh_from_db()
    assert optimized.output_language == "nl"
    assert optimized.optimized_resume.locale == "nl"
    assert optimized.optimized_resume.versions.first().snapshot["locale"] == "nl"
    assert any("Vaardigheden" in change for change in response.data["changes"])


@pytest.mark.django_db
def test_repeatable_section_update_and_delete_use_stable_id(client):
    resume_id = client.post("/api/resumes/", {"title": "CRUD Resume"}, format="json").data["id"]
    created = client.post(
        f"/api/resumes/{resume_id}/skills/",
        {"skill_name": "React", "category": "Technical"},
        format="json",
    )
    assert created.status_code == 201
    skill_id = created.data["id"]

    updated = client.put(
        f"/api/skills/{skill_id}/",
        {"skill_name": "React.js", "category": "Technical"},
        format="json",
    )
    assert updated.status_code == 200
    assert updated.data["name"] == "React.js"
    assert Skill.objects.filter(id=skill_id, name="React.js").count() == 1

    deleted = client.delete(f"/api/skills/{skill_id}/")
    assert deleted.status_code == 204
    assert not Skill.objects.filter(id=skill_id).exists()


@pytest.mark.django_db
def test_normalized_duplicate_sections_are_rejected(client):
    resume_id = client.post("/api/resumes/", {"title": "Duplicate Resume"}, format="json").data["id"]

    first_skill = client.post(
        f"/api/resumes/{resume_id}/skills/",
        {"skill_name": "React", "category": "Technical"},
        format="json",
    )
    duplicate_skill = client.post(
        f"/api/resumes/{resume_id}/skills/",
        {"skill_name": "  REACT  ", "category": "Tool"},
        format="json",
    )
    assert first_skill.status_code == 201
    assert duplicate_skill.status_code == 400
    assert "already added" in str(duplicate_skill.data)
    assert Skill.objects.filter(resume_id=resume_id).count() == 1

    education = {
        "institution": "University of Amsterdam",
        "degree": "Bachelor of Science",
        "field_of_study": "Computer Science",
        "start_date": "2020-09-01",
    }
    assert client.post(f"/api/resumes/{resume_id}/education/", education, format="json").status_code == 201
    duplicate_education = client.post(
        f"/api/resumes/{resume_id}/education/",
        {**education, "institution": " university  OF amsterdam "},
        format="json",
    )
    assert duplicate_education.status_code == 400
    assert Education.objects.filter(resume_id=resume_id).count() == 1

    experience = {
        "company": "GuideWisey",
        "job_title": "Software Engineer",
        "start_date": "2024-01-01",
    }
    assert client.post(f"/api/resumes/{resume_id}/experiences/", experience, format="json").status_code == 201
    duplicate_experience = client.post(
        f"/api/resumes/{resume_id}/experiences/",
        {**experience, "company": " guidewisey ", "job_title": "SOFTWARE ENGINEER"},
        format="json",
    )
    assert duplicate_experience.status_code == 400
    assert WorkExperience.objects.filter(resume_id=resume_id).count() == 1

    project = {"project_name": "Career Suite", "role": "Lead", "description": "Built the product."}
    assert client.post(f"/api/resumes/{resume_id}/projects/", project, format="json").status_code == 201
    duplicate_project = client.post(
        f"/api/resumes/{resume_id}/projects/",
        {**project, "project_name": " career suite ", "role": "LEAD"},
        format="json",
    )
    assert duplicate_project.status_code == 400
    assert Project.objects.filter(resume_id=resume_id).count() == 1


@pytest.mark.django_db
def test_dated_items_create_json_safe_versions_and_do_not_break_later_saves(client):
    resume_id = client.post("/api/resumes/", {"title": "Dated Resume"}, format="json").data["id"]
    experience = client.post(
        f"/api/resumes/{resume_id}/experiences/",
        {
            "company": "GuideWisey",
            "job_title": "Engineer",
            "start_date": "2024-01-15",
            "end_date": "2025-06-30",
            "description": "Built APIs.",
        },
        format="json",
    )
    assert experience.status_code == 201

    skill = client.post(
        f"/api/resumes/{resume_id}/skills/",
        {"skill_name": "Django", "category": "Technical"},
        format="json",
    )
    assert skill.status_code == 201

    resume = Resume.objects.get(id=resume_id)
    dated_snapshot = resume.versions.first()
    assert dated_snapshot.snapshot["experiences"][0]["start_date"] == "2024-01-15"
    assert dated_snapshot.snapshot["experiences"][0]["end_date"] == "2025-06-30"


@pytest.mark.django_db
def test_section_validation_requires_requested_fields(client):
    resume_id = client.post("/api/resumes/", {"title": "Validation Resume"}, format="json").data["id"]
    assert client.post(f"/api/resumes/{resume_id}/skills/", {"skill_name": "React"}, format="json").status_code == 400
    assert (
        client.post(
            f"/api/resumes/{resume_id}/education/",
            {"institution": "University"},
            format="json",
        ).status_code
        == 400
    )
    assert (
        client.post(
            f"/api/resumes/{resume_id}/experiences/",
            {"company": "GuideWisey", "job_title": "Engineer"},
            format="json",
        ).status_code
        == 400
    )


def make_png():
    image = Image.new("RGB", (32, 32), "navy")
    output = io.BytesIO()
    image.save(output, format="PNG")
    return output.getvalue()


def template_defaults(supports_photo=False):
    return {
        "primary_color": "#2563EB",
        "font_family": "Inter",
        "font_size": "medium",
        "spacing": "normal",
        "include_photo": supports_photo,
        "section_order": ["summary", "experience", "education", "skills", "languages"],
    }


def seed_template_catalog():
    templates = [
        ("classic-ats", "Classic ATS", "ATS Friendly", "resumes/classic_ats.html", False, True),
        (
            "modern-professional",
            "Modern Professional",
            "Modern",
            "resumes/modern_professional.html",
            True,
            True,
        ),
        ("executive-clean", "Executive Clean", "Executive", "resumes/executive_clean.html", False, True),
        ("developer-tech", "Developer Tech", "Developer / Tech", "resumes/developer_tech.html", False, True),
        ("student-fresher", "Student Fresher", "Student / Fresher", "resumes/student_fresher.html", True, True),
        ("creative-photo", "Creative Photo", "Photo CV", "resumes/creative_photo.html", True, False),
        ("european-cv", "European CV", "European CV", "resumes/european_cv.html", True, True),
        ("minimal-one-page", "Minimal One Page", "Minimal", "resumes/minimal_one_page.html", False, True),
        ("elegant-timeline", "Elegant Timeline", "Creative", "resumes/elegant_timeline.html", True, False),
        ("compact-recruiter", "Compact Recruiter", "ATS Friendly", "resumes/compact_recruiter.html", False, True),
    ]
    created = []
    for order, (slug, name, category, html_template, supports_photo, ats) in enumerate(templates, 1):
        template, _ = ResumeTemplate.objects.update_or_create(
            slug=slug,
            defaults={
                "name": name,
                "category": category,
                "html_template": html_template,
                "supports_photo": supports_photo,
                "is_ats_friendly": ats,
                "sort_order": order,
                "supported_formats": ["pdf", "docx"],
                "default_settings": template_defaults(supports_photo),
            },
        )
        created.append(template)
    return created


def complete_resume(client, resume):
    client.put(
        f"/api/resumes/{resume.id}/personal/",
        {
            "first_name": "Ada",
            "last_name": "Lovelace",
            "email": "ada@example.com",
            "phone": "123456789",
            "city": "London",
            "country": "UK",
            "professional_title": "Engineer",
        },
        format="json",
    )
    client.put(f"/api/resumes/{resume.id}/summary/", {"text": "Engineering leader."}, format="json")
    client.post(
        f"/api/resumes/{resume.id}/education/",
        {"institution": "University", "degree": "Computer Science"},
        format="json",
    )
    for name in ("Python", "Django", "SQL"):
        client.post(
            f"/api/resumes/{resume.id}/skills/",
            {"skill_name": name, "category": "Technical"},
            format="json",
        )
    client.post(
        f"/api/resumes/{resume.id}/languages/",
        {"name": "English", "proficiency": "Fluent"},
        format="json",
    )


@pytest.mark.django_db
def test_template_catalog_list_and_detail(client):
    templates = seed_template_catalog()
    response = client.get("/api/resume-templates/")
    assert response.status_code == 200
    data = response.data["results"] if isinstance(response.data, dict) else response.data
    assert len(data) == 10
    assert data[0]["name"] == "Classic ATS"
    detail = client.get(f"/api/resume-templates/{templates[1].id}/")
    assert detail.status_code == 200
    assert detail.data["name"] == "Modern Professional"


@pytest.mark.django_db
def test_select_template_and_preview(client, user):
    templates = seed_template_catalog()
    template = templates[1]
    resume = Resume.objects.create(user=user, title="Template Resume")
    invalid = client.post(
        f"/api/resumes/{resume.id}/select-template/",
        {"template_id": templates[0].id, "template_settings": {"include_photo": True}},
        format="json",
    )
    assert invalid.status_code == 400
    selected = client.post(
        f"/api/resumes/{resume.id}/select-template/",
        {
            "template_id": template.id,
            "template_settings": {
                "primary_color": "#123ABC",
                "font_family": "Inter",
                "font_size": "large",
                "spacing": "relaxed",
                "include_photo": True,
                "section_order": ["summary", "skills", "experience", "education"],
            },
        },
        format="json",
    )
    assert selected.status_code == 200
    assert selected.data["template"]["name"] == "Modern Professional"
    assert selected.data["template_settings"]["primary_color"] == "#123ABC"
    resume.refresh_from_db()
    assert resume.template_id == template.id
    version_count = resume.versions.count()
    repeated = client.post(
        f"/api/resumes/{resume.id}/select-template/",
        {
            "template_id": template.id,
            "template_settings": {
                "primary_color": "#123ABC",
                "font_family": "Inter",
                "font_size": "large",
                "spacing": "relaxed",
                "include_photo": True,
                "section_order": ["summary", "skills", "experience", "education"],
            },
        },
        format="json",
    )
    assert repeated.status_code == 200
    assert resume.versions.count() == version_count

    preview = client.post(
        f"/api/resumes/{resume.id}/preview/",
        {"template_id": template.id, "template_settings": {"primary_color": "#654321"}},
        format="json",
    )
    assert preview.status_code == 200
    assert "--accent: #654321" in preview.data["html"]


@pytest.mark.django_db
def test_export_requires_selected_template(client, user):
    resume = Resume.objects.create(user=user, title="No Template")
    response = client.post(f"/api/resumes/{resume.id}/export/pdf/", {}, format="json")
    assert response.status_code == 400
    assert response.data["error"] == "Please select a resume template before generating your CV."


@pytest.mark.django_db
def test_export_works_after_template_selection(client, user):
    template = seed_template_catalog()[0]
    resume = Resume.objects.create(user=user, title="Selected Template")
    complete_resume(client, resume)
    selected = client.post(
        f"/api/resumes/{resume.id}/select-template/",
        {"template_id": template.id, "template_settings": {}},
        format="json",
    )
    assert selected.status_code == 200
    response = client.post(f"/api/resumes/{resume.id}/export/docx/", {}, format="json")
    assert response.status_code == 201


@pytest.mark.django_db
def test_profile_photo_upload_preview_remove_and_owner_scope(client, user, other_user):
    resume = Resume.objects.create(user=user, title="Photo Resume")
    upload = SimpleUploadedFile("profile.png", make_png(), content_type="image/png")
    response = client.post(f"/api/resumes/{resume.id}/photo/upload/", {"photo": upload}, format="multipart")
    assert response.status_code == 201
    assert response.data["has_photo"] is True
    assert response.data["photo_url"].endswith(f"/api/resumes/{resume.id}/photo/")

    preview = client.get(f"/api/resumes/{resume.id}/photo/")
    assert preview.status_code == 200
    assert preview["Content-Type"] == "image/png"

    other_client = APIClient()
    other_client.force_authenticate(other_user)
    assert other_client.get(f"/api/resumes/{resume.id}/photo/").status_code == 404
    assert client.delete(f"/api/resumes/{resume.id}/photo/").status_code == 204
    personal = PersonalDetail.objects.get(resume=resume)
    assert not personal.profile_photo_id
    assert personal.include_photo is False


@pytest.mark.django_db
def test_preview_photo_obeys_template_support_and_include_setting(user):
    templates = seed_template_catalog()
    no_photo_template = templates[0]
    photo_template = templates[1]
    photo = UserFile.objects.create(
        user=user,
        filename="profile.png",
        content_type="image/png",
        file_size=len(make_png()),
        file_data=make_png(),
        purpose="resume_profile_photo",
    )
    resume = Resume.objects.create(
        user=user,
        title="Photo Preview",
        template=photo_template,
        include_photo=True,
        template_settings={**template_defaults(True), "include_photo": True},
    )
    PersonalDetail.objects.create(
        resume=resume,
        profile_photo=photo,
        include_photo=True,
        first_name="Ada",
        last_name="Lovelace",
    )
    assert "data:image/png;base64," in render_resume_html(resume)

    resume.include_photo = False
    resume.save(update_fields=["include_photo"])
    assert "data:image/png;base64," not in render_resume_html(resume)

    resume.include_photo = True
    resume.template = no_photo_template
    resume.save(update_fields=["include_photo", "template"])
    assert "data:image/png;base64," not in render_resume_html(resume)


@pytest.mark.django_db
def test_docx_photo_respects_include_toggle(user):
    template = ResumeTemplate.objects.create(slug="photo", name="Photo", supports_photo=True)
    resume = Resume.objects.create(user=user, title="Photo Export", template=template)
    photo = UserFile.objects.create(
        user=user,
        filename="profile.png",
        content_type="image/png",
        file_size=len(make_png()),
        file_data=make_png(),
        purpose="resume_profile_photo",
    )
    personal = PersonalDetail.objects.create(
        resume=resume,
        profile_photo=photo,
        include_photo=True,
        first_name="Ada",
        last_name="Lovelace",
        email="ada@example.com",
        phone="123",
        city="London",
        country="UK",
        professional_title="Engineer",
    )
    client = APIClient()
    client.force_authenticate(user)
    client.put(f"/api/resumes/{resume.id}/summary/", {"text": "Engineer."}, format="json")
    client.post(
        f"/api/resumes/{resume.id}/education/",
        {"institution": "University", "degree": "BSc"},
        format="json",
    )
    for name in ("Python", "Django", "SQL"):
        client.post(
            f"/api/resumes/{resume.id}/skills/",
            {"skill_name": name, "category": "Technical"},
            format="json",
        )
    client.post(
        f"/api/resumes/{resume.id}/languages/",
        {"name": "English", "proficiency": "Fluent"},
        format="json",
    )

    assert get_export_photo(resume).id == photo.id
    with zipfile.ZipFile(io.BytesIO(generate_docx(resume))) as archive:
        assert any(name.startswith("word/media/") for name in archive.namelist())

    personal.include_photo = False
    personal.save(update_fields=["include_photo"])
    assert get_export_photo(resume) is None
    with zipfile.ZipFile(io.BytesIO(generate_docx(resume))) as archive:
        assert not any(name.startswith("word/media/") for name in archive.namelist())

    personal.include_photo = True
    personal.save(update_fields=["include_photo"])
    template.supports_photo = False
    template.save(update_fields=["supports_photo"])
    assert get_export_photo(resume) is None


@pytest.mark.django_db
def test_job_url_endpoint_success_and_graceful_failure(client):
    parsed = {
        "title": "Backend Engineer",
        "company": "GuideWisey",
        "location": "Amsterdam",
        "raw_text": "Required skills: Python, Docker",
        "required_skills": ["Python", "Docker"],
        "preferred_skills": [],
        "responsibilities": [],
        "education": [],
        "certifications": [],
        "keywords": ["python", "docker"],
        "skills": ["python", "docker"],
    }
    with patch("apps.jobs.views.parse_job_url", return_value=parsed):
        response = client.post("/api/jobs/parse-url/", {"url": "https://example.com/job"}, format="json")
    assert response.status_code == 201
    assert response.data["success"] is True
    assert response.data["raw_text"] == parsed["raw_text"]
    assert response.data["parsed_json"]["required_skills"] == ["Python", "Docker"]

    with patch("apps.jobs.views.parse_job_url", side_effect=requests.Timeout("blocked")):
        failure = client.post("/api/jobs/parse-url/", {"url": "https://example.com/blocked"}, format="json")
    assert failure.status_code == 400
    assert failure.data["success"] is False
    assert "paste the job description manually" in failure.data["error"]


@pytest.mark.django_db
def test_job_url_endpoint_persists_long_tracking_url(client):
    long_url = "https://jobs.example.com/customer-support?" + "&".join(
        f"tracking_parameter_{index}={'x' * 40}" for index in range(8)
    )
    assert len(long_url) > 200
    parsed = {
        "title": "Customer Support Specialist",
        "job_title": "Customer Support Specialist",
        "company": "Example",
        "location": "Amsterdam",
        "raw_text": "Required skills: Customer Service",
        "required_skills": ["Customer Service"],
        "preferred_skills": [],
        "responsibilities": [],
        "education": [],
        "education_requirements": [],
        "certifications": [],
        "tools": [],
        "technologies": [],
        "keywords": ["customer", "service"],
        "language_requirements": [],
    }

    with patch("apps.jobs.views.parse_job_url", return_value=parsed):
        response = client.post(
            "/api/jobs/parse-url/",
            {"url": long_url, "language": "en"},
            format="json",
        )

    assert response.status_code == 201
    job = JobDescription.objects.get(id=response.data["job_description"]["id"])
    temporary = TemporaryJobDescription.objects.get(id=response.data["temporary_id"])
    assert job.source_url == long_url
    assert temporary.source_url == long_url


@pytest.mark.django_db
def test_job_url_endpoint_rejects_page_without_readable_text(client):
    with patch("apps.jobs.services._validate_public_url"), patch("apps.jobs.services.requests.get") as get:
        get.return_value.is_redirect = False
        get.return_value.is_permanent_redirect = False
        get.return_value.content = b""
        get.return_value.text = ""
        get.return_value.raise_for_status.return_value = None
        response = client.post("/api/jobs/parse-url/", {"url": "https://example.com/empty"}, format="json")

    assert response.status_code == 400
    assert response.data["success"] is False
    assert "paste the job description manually" in response.data["error"]


@pytest.mark.django_db
def test_autocomplete_returns_curated_and_user_values(client, user):
    resume = Resume.objects.create(user=user, title="Autocomplete")
    Skill.objects.create(resume=resume, name="React Native", category="Technical")
    WorkExperience.objects.create(
        resume=resume,
        employer="GuideWisey",
        job_title="Platform Engineer",
        start_date="2024-01-01",
    )
    Education.objects.create(
        resume=resume,
        institution="Amsterdam Tech Institute",
        degree="Bachelor of Engineering",
    )

    skills = client.get("/api/autocomplete/skills/?q=react")
    assert skills.status_code == 200
    assert {"React", "React Native"}.issubset({item["value"] for item in skills.data["results"]})
    titles = client.get("/api/autocomplete/job-titles/?q=platform")
    assert titles.data["results"][0]["value"] == "Platform Engineer"
    companies = client.get("/api/autocomplete/companies/?q=guide")
    assert companies.data["results"][0]["value"] == "GuideWisey"
    schools = client.get("/api/autocomplete/schools/?q=amsterdam")
    assert any(item["value"] == "Amsterdam Tech Institute" for item in schools.data["results"])
    degrees = client.get("/api/autocomplete/degrees/?q=bachelor")
    assert any(item["value"] == "Bachelor of Engineering" for item in degrees.data["results"])
    locations = client.get("/api/autocomplete/locations/?q=am")
    assert locations.status_code == 200
    assert locations.data["results"]
    dutch_skills = client.get("/api/autocomplete/skills/?q=commu&lang=nl")
    assert dutch_skills.data["results"][0]["label"] == "Communicatie"
    assert dutch_skills.data["results"][0]["value"] == "Communication"


@pytest.mark.django_db
def test_optimized_resume_clone_deduplicates_legacy_snapshot_skills(user):
    source = Resume.objects.create(user=user, title="Legacy Resume")
    snapshot = {
        "personal": {},
        "summary": "",
        "skills": [
            {"name": "Java Script", "category": "Technical", "position": 0},
            {"name": " java   script ", "category": "Technical", "position": 1},
        ],
        "education": [],
        "experiences": [],
        "projects": [],
        "certifications": [],
        "languages": [],
        "awards": [],
        "references": [],
    }

    optimized, _ = create_resume_from_snapshot(user, source, snapshot)

    assert list(optimized.skill_set.values_list("name", flat=True)) == ["Java Script"]


@pytest.mark.django_db
def test_optimize_target_and_declined_skills(client):
    resume_id = client.post("/api/resumes/", {"title": "Target Resume"}, format="json").data["id"]
    template = ResumeTemplate.objects.create(
        slug="target-template",
        name="Target Template",
        html_template="resumes/classic_ats.html",
        default_settings=template_defaults(),
    )
    source_resume = Resume.objects.get(id=resume_id)
    source_resume.template = template
    source_resume.template_settings = {**template_defaults(), "primary_color": "#123ABC"}
    source_resume.save(update_fields=["template", "template_settings"])
    client.put(
        f"/api/resumes/{resume_id}/personal/",
        {"professional_title": "Engineer"},
        format="json",
    )
    job = client.post(
        "/api/jobs/parse-text/",
        {"title": "Platform Engineer", "text": "Required skills: Docker, AWS"},
        format="json",
    ).data["job_description"]
    match = client.post(
        "/api/job-match/analyze/",
        {"resume_id": resume_id, "job_description_id": job["id"]},
        format="json",
    ).data

    invalid = client.post(
        f"/api/job-match/{match['id']}/optimize/",
        {"target_score": 101},
        format="json",
    )
    assert invalid.status_code == 400

    optimized = client.post(
        f"/api/job-match/{match['id']}/optimize/",
        {
            "target_score": 95,
            "confirmed_skills": [{"skill": "Docker", "confirmed": True, "evidence": "Used for deployment"}],
            "declined_skills": ["AWS"],
            "output_language": "en",
        },
        format="json",
    )
    assert optimized.status_code == 201
    assert optimized.data["optimized_resume_record"]["id"] == optimized.data["optimized_resume_id"]
    assert optimized.data["optimized_resume_record"]["selected_template"]["id"] == template.id
    optimized_resume = Resume.objects.get(id=optimized.data["optimized_resume_id"])
    assert optimized_resume.template_id == template.id
    assert optimized_resume.template_settings["primary_color"] == "#123ABC"
    names = set(optimized_resume.skill_set.values_list("name", flat=True))
    assert "Docker" in names
    assert "AWS" not in names
    assert "Used for deployment" in optimized_resume.summary.text
    assert optimized.data["target_score"] == 95
    assert any("declined" in gap for gap in optimized.data["remaining_gaps"])


@pytest.mark.django_db
def test_optimize_does_not_add_confirmed_skill_unrelated_to_job(client):
    resume_id = client.post("/api/resumes/", {"title": "Honest Resume"}, format="json").data["id"]
    job = client.post(
        "/api/jobs/parse-text/",
        {"title": "Platform Engineer", "text": "Required skills: Docker"},
        format="json",
    ).data["job_description"]
    match = client.post(
        "/api/job-match/analyze/",
        {"resume_id": resume_id, "job_description_id": job["id"]},
        format="json",
    ).data

    optimized = client.post(
        f"/api/job-match/{match['id']}/optimize/",
        {
            "target_score": 90,
            "confirmed_skills": [{"skill": "AWS", "confirmed": True, "evidence": "Unrelated claim"}],
        },
        format="json",
    )
    assert optimized.status_code == 201
    optimized_resume = Resume.objects.get(id=optimized.data["optimized_resume_id"])
    assert not optimized_resume.skill_set.filter(name__iexact="AWS").exists()


@pytest.mark.django_db
def test_auto_fill_from_job_creates_reviewable_starter_without_fabricating_credentials(client):
    response = client.post(
        "/api/resume-builder/auto-fill-from-job/",
        {
            "job_description_text": (
                "Job title: Customer Support Specialist\n"
                "Required skills: Customer service, Salesforce, communication\n"
                "Responsibilities: Help customers; resolve support requests\n"
                "Education: Bachelor degree\n"
                "Certification: Salesforce certification"
            ),
            "target_language": "en",
        },
        format="json",
    )

    assert response.status_code == 201
    resume = Resume.objects.get(id=response.data["resume_id"])
    payload = response.data["builder_payload"]
    suggested = {item["name"]: item["status"] for item in payload["skills"]}

    assert resume.personal.professional_title == "Customer Support Specialist"
    assert not WorkExperience.objects.filter(resume=resume).exists()
    assert not resume.education_set.exists()
    assert not resume.certification_set.exists()
    assert not resume.skill_set.filter(name__iexact="Salesforce").exists()
    assert suggested["Salesforce"] == "needs_confirmation"
    assert payload["experience"][0]["status"] == "needs_user_input"
    assert payload["education"][0]["status"] == "needs_user_input"
    assert payload["summary"]["status"] == "needs_user_input"
    assert any(item["type"] == "skill_confirmation" for item in payload["review_questions"])
    assert TemporaryGeneratedResume.objects.filter(id=response.data["draft_id"]).exists()
    reloaded = client.get(f"/api/resume-builder/auto-fill-drafts/{response.data['draft_id']}/")
    assert reloaded.status_code == 200
    assert reloaded.data["builder_payload"] == payload


@pytest.mark.django_db
def test_auto_fill_from_existing_resume_preserves_facts_and_only_suggests_missing_skills(client, user):
    source = Resume.objects.create(user=user, title="Source Resume", locale="en")
    PersonalDetail.objects.create(
        resume=source,
        first_name="Ada",
        last_name="Lovelace",
        email="ada@example.com",
        professional_title="Support Agent",
    )
    ResumeSummary.objects.create(resume=source, text="Supported enterprise customers.")
    WorkExperience.objects.create(
        resume=source,
        employer="Real Company",
        job_title="Support Agent",
        start_date="2022-01-01",
        description="Resolved customer requests. Documented recurring issues.",
    )
    Skill.objects.create(resume=source, name="Customer Service", category="Functional")

    response = client.post(
        "/api/resume-builder/auto-fill-from-job/",
        {
            "job_description_text": (
                "Job title: Customer Support Specialist\n"
                "Required skills: Customer service, Salesforce\n"
                "Responsibilities: Resolve customer requests"
            ),
            "resume_id": source.id,
            "target_language": "en",
        },
        format="json",
    )

    assert response.status_code == 201
    generated = Resume.objects.get(id=response.data["resume_id"])
    experience = WorkExperience.objects.get(resume=generated)
    skills = {item["name"]: item["status"] for item in response.data["builder_payload"]["skills"]}

    assert generated.id != source.id
    assert generated.personal.first_name == "Ada"
    assert experience.employer == "Real Company"
    assert experience.job_title == "Support Agent"
    assert str(experience.start_date) == "2022-01-01"
    assert "Resolved customer requests." in experience.description
    assert skills["Customer Service"] == "confirmed"
    assert skills["Salesforce"] == "needs_confirmation"
    assert not generated.skill_set.filter(name__iexact="Salesforce").exists()
    assert generated.versions.filter(source="auto_fill_from_job").exists()


@pytest.mark.django_db
def test_auto_fill_from_parsed_upload_uses_only_uploaded_facts(client, user):
    upload = ResumeUpload.objects.create(
        user=user,
        filename="resume.txt",
        content_type="text/plain",
        file_size=10,
        file_data=b"resume",
        status="completed",
        parsed_json={
            "personal": {"first_name": "Grace", "email": "grace@example.com"},
            "summary": "Experienced in customer communication.",
            "skills": [{"name": "Communication"}],
            "parser_version": 2,
        },
    )

    response = client.post(
        "/api/resume-builder/auto-fill-from-job/",
        {
            "job_description_text": (
                "Functie: Klantenservice Medewerker\n"
                "Vereiste vaardigheden: Communicatie, Salesforce\n"
                "Verantwoordelijkheden: Klanten helpen"
            ),
            "uploaded_resume_id": str(upload.id),
            "target_language": "nl",
        },
        format="json",
    )

    assert response.status_code == 201
    resume = Resume.objects.get(id=response.data["resume_id"])
    statuses = {item["name"]: item["status"] for item in response.data["builder_payload"]["skills"]}

    assert resume.locale == "nl"
    assert resume.personal.first_name == "Grace"
    assert resume.personal.email == "grace@example.com"
    assert resume.skill_set.filter(name="Communication").exists()
    assert not resume.skill_set.filter(name="Salesforce").exists()
    assert statuses["Communication"] == "confirmed"
    assert statuses["Salesforce"] == "needs_confirmation"
    assert any("controleer" in warning.lower() for warning in response.data["warnings"])


@pytest.mark.django_db
def test_auto_fill_job_url_uses_secure_existing_parser(client):
    parsed = {
        "title": "Backend Engineer",
        "job_title": "Backend Engineer",
        "company": "GuideWisey",
        "location": "Amsterdam",
        "seniority": "Senior",
        "required_skills": ["Docker"],
        "preferred_skills": [],
        "responsibilities": ["Build APIs"],
        "tools": ["Docker"],
        "technologies": ["Docker"],
        "education_requirements": [],
        "certifications": [],
        "keywords": ["docker", "apis"],
        "language_requirements": ["English"],
        "raw_text": "Backend Engineer role",
    }
    with patch("apps.resumes.auto_fill.parse_job_url", return_value=parsed) as parser:
        response = client.post(
            "/api/resume-builder/auto-fill-from-job/",
            {"job_description_url": "https://example.com/jobs/backend", "target_language": "en"},
            format="json",
        )

    assert response.status_code == 201
    parser.assert_called_once_with("https://example.com/jobs/backend")
    assert response.data["builder_payload"]["job_description"]["company"] == "GuideWisey"
